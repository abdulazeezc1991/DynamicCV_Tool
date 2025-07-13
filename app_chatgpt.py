import os
import docx
from flask import Flask, render_template, request, send_file
import openai
from io import BytesIO

app = Flask(__name__)
openai.api_key = "Your Api Key"  # <-- Replace this key

def extract_text_from_docx(file):
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def generate_tailored_cv(cv_text, jd_text):
    messages = [
        {
            "role": "user",
            "content": f"""
Please carefully read and understand my existing resume below, including my work history, skills, and experience. DO NOT rewrite, respond, or provide any feedback now. Just learn and remember this information for my future requests where I will ask you to help rewrite a cleaner, shorter resume based on a specific job description. Here is my existing resume:
{cv_text}
"""
        },
        {
            "role": "user",
            "content": f"""
Write a 2–3 line resume summary using the X-Y-Z method. Mention total years of experience, key technical skills (3 skills) from the job description, and domain exposure (e.g., finance, government, healthcare, etc.). Highlight 1–2 quantifiable achievements with numbers and include the current company name or “last company.” Keep it under 50 words, ATS-friendly by using keywords from the job description. Avoid buzzwords and formal or robotic phrases. Make it sound personal, confident, and genuine. Add a separate line stating “Open to relocation.”
Job Description:
{jd_text}
"""
        },
        {
            "role": "user",
            "content": f"""
For each previous job role, write detailed bullet points that reflect my key responsibilities and achievements. Include job title, company name, location, and dates of employment in reverse chronological order. Use 5-6 bullet points for my latest job and 3-4 for earlier roles. Include quantifiable results where possible. Write in a natural, human tone—not like AI or robotic text. Use relevant keywords from the job description provided below.
Job Description:
{jd_text}
"""
        },
        {
            "role": "user",
            "content": f"""
Create a skills section listing technical skills and relevant soft skills from the job description. Avoid generic terms, buzzwords, and clichés. Provide only the skill headings—no descriptions.
Job Description:
{jd_text}
"""
        },
        {
            "role": "user",
            "content": """
List my education in reverse chronological order, including degree, institution name, location, and start and end dates (month/year).
"""
        },
        {
            "role": "user",
            "content": """
List relevant certifications with issuing organizations and dates. If I don’t have certificates but completed training, write '[Certification Name] trained' instead.
"""
        },
        {
            "role": "user",
            "content": f"""
Write a resume in 500-600 words using the format below. Use the Job Description to choose relevant keywords and highlight achievements with numbers. Keep the tone natural, confident, and professional.

Sections:
1. Contact Information
- Full name, phone number, email, LinkedIn, location

2. Professional Summary (Max 50 words)
- Total years of experience in relevant domain
- Skilled in top 3 skills from the job description
- Delivered quantifiable results at last company
- Add: “Open to relocation.”

3. Work Experience (Max 350 words)
- Latest job: 5-6 bullet points with results and keywords
- Earlier jobs: 3-4 bullet points

4. Skills (Max 50 words)
- Technical and soft skills relevant to the JD

5. Education (Max 50 words)
- Degree, institution, location, dates

6. Certifications and Training (Max 50 words)
- Certification name, org, date or “trained”

7. Additional Sections (Optional, Max 50 words)
- Awards, projects, volunteer work, if relevant

Guidelines:
- Use simple, human, confident language
- ATS-friendly formatting
- Avoid exaggeration and filler content
- Check grammar and tone

Job Description:
{jd_text}

Resume:
{cv_text}
"""
        }
    ]

    response = openai.ChatCompletion.create(
        model="gpt-3.5-turbo",
        messages=messages,
        temperature=0.7
    )
    return response.choices[0].message.content.strip()


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        cv_file = request.files["cv"]
        jd_text = request.form["jd"]

        if cv_file and jd_text:
            cv_text = extract_text_from_docx(cv_file)
            tailored_cv = generate_tailored_cv(cv_text, jd_text)

            # Create downloadable docx
            doc = docx.Document()
            for para in tailored_cv.split("\n"):
                doc.add_paragraph(para)
            file_stream = BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)

            return send_file(file_stream, as_attachment=True, download_name="tailored_cv.docx")
    return render_template("index.html")

if __name__ == "__main__":
    app.run(debug=True)