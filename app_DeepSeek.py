import os
import docx
from flask import Flask, render_template, request, send_file
import requests
from io import BytesIO

app = Flask(__name__)
DEEPSEEK_API_KEY = os.environ.get("DEEPSEEK_API_KEY")
DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"

def extract_text_from_docx(file):
    doc = docx.Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def generate_tailored_cv(cv_text, jd_text):
    messages = [
        # ... (keep all your existing message blocks exactly as they are) ...
    ]
    
    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json"
    }
    
    payload = {
        "model": "deepseek-chat",  # Confirm this is the correct model name
        "messages": [
            {"role": "system", "content": "You are a professional resume assistant."},
            {"role": "user", "content": f"Resume: {cv_text}\n\nJob Description: {jd_text}"}
        ],
        "temperature": 0.7,
        "max_tokens": 2000  # Add if needed
    }
    
    try:
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=payload)
        response.raise_for_status()
        data = response.json()
        return data['choices'][0]['message']['content'].strip()
    except Exception as e:
        print(f"Error: {e}")
        return "Error generating CV. Please check logs."

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        cv_file = request.files["cv"]
        jd_text = request.form["jd"]

        if cv_file and jd_text:
            cv_text = extract_text_from_docx(cv_file)
            tailored_cv = generate_tailored_cv(cv_text, jd_text)

            doc = docx.Document()
            for para in tailored_cv.split("\n"):
                doc.add_paragraph(para)
            file_stream = BytesIO()
            doc.save(file_stream)
            file_stream.seek(0)

            return send_file(file_stream, as_attachment=True, download_name="tailored_cv.docx")
    return render_template("index.html")

if __name__ == "__main__":
    app.run(debug=True)